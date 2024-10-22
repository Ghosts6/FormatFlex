import fitz
import docx  
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from xml.etree import ElementTree as ET

def pdf_to_word(pdf_path, word_path):
    pdf_document = fitz.open(pdf_path)
    doc = Document()
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        text = page.get_text("text")
        doc.add_paragraph(text)

    doc.save(word_path)
    print(f"PDF converted to Word: {word_path}")

def xml_to_pdf(file_path, file_2_path):
    with open(file_path, 'r') as xml_file:
        xml_data = xml_file.read()
    
    root = ET.fromstring(xml_data)
    
    c = canvas.Canvas(file_2_path, pagesize=letter)
    
    y_offset = 750
    for element in root:
        if element.tag == 'title':
            c.setFont("Helvetica-Bold", 16)
            c.drawString(100, y_offset, element.text)
            y_offset -= 20
            
        elif element.tag == 'paragraph':
            c.setFont("Helvetica", 12)
            c.drawString(100, y_offset, element.text)
            y_offset -= 15
            
    c.save()
    print(f"XML converted to PDF: {file_2_path}")

def xml_to_pdf(file_path, file_2_path):
    with open(file_path, 'r') as xml_file:
        xml_data = xml_file.read()
    
    root = ET.fromstring(xml_data)
    
    c = canvas.Canvas(file_2_path, pagesize=letter)
    
    y_offset = 750
    for element in root:
        if element.tag == 'title':
            c.setFont("Helvetica-Bold", 16)
            c.drawString(100, y_offset, element.text)
            y_offset -= 20
            
        elif element.tag == 'paragraph':
            c.setFont("Helvetica", 12)
            c.drawString(100, y_offset, element.text)
            y_offset -= 15
            
    c.save()
    print(f"XML converted to PDF: {file_2_path}")

def xml_to_word(file_path, file_2_path):
    with open(file_path, 'r') as xml_file:
        xml_data = xml_file.read()
    
    root = ET.fromstring(xml_data)
    
    doc = Document()
    
    for element in root:
        if element.tag == 'title':
            doc.add_heading(element.text, level=1)
            
        elif element.tag == 'paragraph':
            doc.add_paragraph(element.text)
    
    doc.save(file_2_path)
    print(f"XML converted to Word: {file_2_path}")

def xml_to_txt(file_path, file_2_path):
    with open(file_path, 'r') as xml_file:
        xml_data = xml_file.read()
    
    root = ET.fromstring(xml_data)
    
    with open(file_2_path, 'w') as txt_file:
        for element in root:
            txt_file.write(element.text + '\n')
    
    print(f"XML converted to TXT: {file_2_path}")

def pdf_to_xml(pdf_path, xml_path):
    pdf_document = fitz.open(pdf_path)
    
    root = ET.Element("document")
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        text = page.get_text()
        
        page_element = ET.SubElement(root, "page")
        
        page_element.text = text
    
    tree = ET.ElementTree(root)
    tree.write(xml_path)
    
    print(f"PDF converted to XML: {xml_path}")

def pdf_to_txt(pdf_path, txt_path):
    pdf_document = fitz.open(pdf_path)
    text = ""
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        text += page.get_text()
    
    with open(txt_path, 'w') as txt_file:
        txt_file.write(text)
    
    print(f"PDF converted to TXT: {txt_path}")

def word_to_xml(word_path, xml_path):
    doc = docx.Document(word_path)
    
    root = ET.Element("document")
    
    for paragraph in doc.paragraphs:
        paragraph_element = ET.SubElement(root, "paragraph")

        paragraph_element.text = paragraph.text

    tree = ET.ElementTree(root)
    tree.write(xml_path)
    
    print(f"Word converted to XML: {xml_path}")

def word_to_pdf(file_path, file_2_path):
    doc = docx.Document(file_path)
    c = canvas.Canvas(file_2_path, pagesize=letter)
    
    for paragraph in doc.paragraphs:
        c.drawString(100, 750, paragraph.text)
        c.showPage()
    
    c.save()
    
    print(f"Word converted to PDF: {file_2_path}")

def word_to_txt(file_path, file_2_path):
    doc = docx.Document(file_path)
    
    with open(file_2_path, 'w') as txt_file:
        for paragraph in doc.paragraphs:
            txt_file.write(paragraph.text + '\n')
    
    print(f"Word converted to TXT: {file_2_path}")

def txt_to_xml(file_path, file_2_path):
    with open(file_path, 'r') as txt_file:
        txt_data = txt_file.read()
    
    root = ET.Element("document")
    
    for line in txt_data.split('\n'):
        paragraph_element = ET.SubElement(root, "paragraph")
        paragraph_element.text = line
    
    tree = ET.ElementTree(root)
    tree.write(file_2_path)
    
    print(f"TXT converted to XML: {file_2_path}")

def txt_to_pdf(file_path, file_2_path):
    with open(file_path, 'r') as txt_file:
        txt_data = txt_file.read()
    
    c = canvas.Canvas(file_2_path, pagesize=letter)
    c.drawString(100, 750, txt_data)
    c.save()
    
    print(f"TXT converted to PDF: {file_2_path}")

def txt_to_word(file_path, file_2_path):
    doc = docx.Document()
    
    with open(file_path, 'r') as txt_file:
        txt_data = txt_file.read()
    
    for line in txt_data.split('\n'):
        doc.add_paragraph(line)
    
    doc.save(file_2_path)
    
    print(f"TXT converted to Word: {file_2_path}")

print("---Welcome to our conversion program---")
choice = input("Select your input file type (xml/pdf/word/txt): ")
convert = input("Select the output file type (xml/pdf/word/txt): ")
choice_path = input("Please enter the path of the input file: ")
convert_path = input("Please enter the path of the output file: ")

if choice == "xml":
    if convert == "pdf":
        xml_to_pdf(choice_path, convert_path)
    elif convert == "word":
        xml_to_word(choice_path, convert_path)
    elif convert == "txt":
        xml_to_txt(choice_path, convert_path)
    else:
        print("Wrong type or same type as input file, please try again.")

elif choice == "pdf":
    if convert == "xml":
        pdf_to_xml(choice_path, convert_path)
    elif convert == "word":
        pdf_to_word(choice_path, convert_path)
    elif convert == "txt":
        pdf_to_txt(choice_path, convert_path)
    else:
        print("Wrong type or same type as input file, please try again.")

elif choice == "word":
    if convert == "pdf":
        word_to_pdf(choice_path, convert_path)
    elif convert == "xml":
        word_to_xml(choice_path, convert_path)
    elif convert == "txt":
        word_to_txt(choice_path, convert_path)
    else:
        print("Wrong type or same type as input file, please try again.")

elif choice == "txt":
    if convert == "xml":
        txt_to_xml(choice_path, convert_path)
    elif convert == "pdf":
        txt_to_pdf(choice_path, convert_path)
    elif convert == "word":
        txt_to_word(choice_path, convert_path)
    else:
        print("Wrong type or same type as input file, please try again.")
else:
    print("Invalid input file type. Please select xml/pdf/word/txt.")
